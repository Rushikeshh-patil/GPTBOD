import openai
import pandas as pd
import docx

#Setting authorization
openai.api_key = ""

#Reading Excel
def import_excel_to_pandas(filepath):
    df = pd.read_excel(filepath)
    return df

#generating the prompt based on the excel sheet data

def generate_prompt():
    df = import_excel_to_pandas('/Users/rushikesh/Library/CloudStorage/GoogleDrive-rdp352@nyu.edu/My Drive/Rushikesh-Windows/GPTBOD/mechanical_list.xlsm')
    system = df[df['CATEGORY'] == 'AIR DISTRIBUTION']
    equipment = system['EQUIPMENT TAG'].tolist()
    
generate_prompt()

#Sending to Openai to generate text
def generate_narrative(query, temperature):
    model_engine = "text-davinci-003" # replace with your preferred GPT-3 model

    prompt = query
    response = openai.Completion.create(
        engine=model_engine,
        prompt=prompt,
        temperature=temperature
    )

    return response.choices[0].text

#save to word file
def create_word_doc(text):
    # create a new Word document
    doc = docx.Document()

    # add a header to the document
    header = doc.sections[0].header
    header.paragraphs[0].text = "Mechanical and plumbing Narrative"
    
    # add a footer to the document with page numbers
    footer = doc.sections[0].footer
    footer_table = footer.add_table(rows=1, cols=1, width=1)
    footer_table.autofit = True
    footer_table.style = 'Table Grid'
    footer_table.rows[0].cells[0].text = "Page {PAGE} of {NUMPAGES}"

    # add the text to the document with specified formatting
    paragraph = doc.add_paragraph(text)
    font = paragraph.style.font
    font.name = 'Calibri'
    font.size = docx.shared.Pt(12)
    font.color.rgb = docx.shared.RGBColor(0, 0, 0)

    # align the text to the left
    paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

    # save the Word document to a file
    doc.save("MP_Narrative.docx")




